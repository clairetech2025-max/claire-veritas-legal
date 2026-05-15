from __future__ import annotations

import math
import json
import re
import time
from collections import Counter
from io import BytesIO
from typing import Any, Dict, List, Optional

try:
    import docx  # type: ignore
except Exception:  # pragma: no cover - optional export dependency
    docx = None


COURT_PROFILES: List[Dict[str, Any]] = [
    {
        "id": "federal_district_civil",
        "name": "Federal District Civil",
        "scope": "National baseline for federal civil litigation",
        "caption_notes": ["Use court caption and docket number", "Verify local rules before filing"],
        "motion_notes": ["Include proposed order where required", "Verify page limits and judge-specific rules"],
        "artifact_defaults": ["complaint", "summons", "civil cover sheet", "motion to compel", "motion to dismiss", "summary judgment", "declaration", "exhibit index"],
        "local_rules_source": "U.S. Courts national rules + local district rules",
        "template_priority": ["motion_to_compel", "motion_to_dismiss", "summary_judgment", "declaration", "exhibit_index"],
        "page_limit_hint": "Vary by court and judge; verify locally",
    },
    {
        "id": "federal_appellate",
        "name": "Federal Appellate",
        "scope": "Circuit appeal workflow",
        "caption_notes": ["Track circuit and case number carefully", "Use appendix and certificate of compliance workflows"],
        "motion_notes": ["Briefs, appendices, and motion practice follow FRAP and circuit rules"],
        "artifact_defaults": ["notice of appeal", "appellant brief", "appellee brief", "reply brief", "appendix", "certificate of compliance"],
        "local_rules_source": "FRAP + circuit local rules",
        "template_priority": ["appendix", "brief", "certificate_of_compliance", "timeline_summary"],
        "page_limit_hint": "Circuit-specific; verify before export",
    },
    {
        "id": "federal_bankruptcy",
        "name": "Federal Bankruptcy",
        "scope": "Bankruptcy case and adversary workflow",
        "caption_notes": ["Use official forms without alteration where required"],
        "motion_notes": ["Bankruptcy form fidelity is strict; verify chapter-specific rules"],
        "artifact_defaults": ["petition", "schedules", "statement of financial affairs", "proof of claim", "motion", "order"],
        "local_rules_source": "FRBP + official bankruptcy forms + local bankruptcy rules",
        "template_priority": ["motion_to_compel", "motion", "declaration", "exhibit_index"],
        "page_limit_hint": "Form-driven; local rules still matter",
    },
    {
        "id": "nd_california",
        "name": "Northern District of California",
        "scope": "Example district profile with explicit motion workflow",
        "caption_notes": ["Use court caption and filing number", "Verify judge-specific instructions"],
        "motion_notes": ["Motion / opposition / reply page limits are court- and judge-driven", "Proposed order often required", "Courtesy copies generally not submitted unless required"],
        "artifact_defaults": ["motion", "opposition", "reply", "proposed order", "exhibit index"],
        "local_rules_source": "N.D. Cal. local rules and judge procedures",
        "template_priority": ["motion_to_compel", "motion_to_dismiss", "summary_judgment", "proposed_order"],
        "page_limit_hint": "Verify local rule and judge standing order",
    },
    {
        "id": "cd_california",
        "name": "Central District of California",
        "scope": "Example district profile with chambers-copy workflow",
        "caption_notes": ["Check e-filing and chambers-copy requirements"],
        "motion_notes": ["Proposed documents may need separate lodging and chambers delivery", "Judge-specific procedures govern final packaging"],
        "artifact_defaults": ["motion", "proposed order", "lodged copy", "exhibit index"],
        "local_rules_source": "C.D. Cal. local rules and procedures",
        "template_priority": ["motion_to_compel", "motion_to_dismiss", "summary_judgment", "declaration"],
        "page_limit_hint": "Verify local rule and judge-specific instructions",
    },
]


FILING_TEMPLATES: List[Dict[str, Any]] = [
    {
        "id": "motion_to_compel",
        "title": "Motion to Compel + MPA",
        "category": "motion",
        "purpose": "Compel deficient discovery responses or production.",
        "sections": ["caption", "notice", "motion", "memorandum of points and authorities", "statement of facts", "argument", "request for relief", "declaration", "exhibit index", "proposed order", "proof of service"],
        "required_inputs": ["discovery requests", "responses or non-responses", "deadlines", "case posture"],
    },
    {
        "id": "motion_to_dismiss",
        "title": "Motion to Dismiss",
        "category": "motion",
        "purpose": "Challenge pleading sufficiency or jurisdictional defects.",
        "sections": ["caption", "notice", "motion", "memorandum of points and authorities", "grounds", "argument", "request for relief", "proposed order", "proof of service"],
        "required_inputs": ["complaint", "causes of action", "jurisdiction facts", "defense theory"],
    },
    {
        "id": "summary_judgment",
        "title": "Motion for Summary Judgment",
        "category": "motion",
        "purpose": "Seek judgment based on undisputed facts and governing law.",
        "sections": ["caption", "notice", "motion", "memorandum of points and authorities", "statement of undisputed facts", "argument", "declaration", "exhibit index", "proposed order", "proof of service"],
        "required_inputs": ["material facts", "depositions", "declarations", "key exhibits"],
    },
    {
        "id": "request_for_admissions",
        "title": "Requests for Admissions",
        "category": "discovery",
        "purpose": "Pin down factual admissions and narrow disputes.",
        "sections": ["caption", "requests", "definitions", "instructions", "service"],
        "required_inputs": ["target facts", "identity of responding party", "discovery deadline"],
    },
    {
        "id": "interrogatories",
        "title": "Interrogatories",
        "category": "discovery",
        "purpose": "Force narrative answers and preserve impeachment material.",
        "sections": ["caption", "interrogatories", "definitions", "instructions", "service"],
        "required_inputs": ["issue list", "custodians", "timeline gaps"],
    },
    {
        "id": "requests_for_production",
        "title": "Requests for Production",
        "category": "discovery",
        "purpose": "Obtain documents, ESI, and metadata.",
        "sections": ["caption", "requests", "definitions", "instructions", "service"],
        "required_inputs": ["document categories", "custodians", "time frame"],
    },
    {
        "id": "declaration",
        "title": "Declaration",
        "category": "supporting",
        "purpose": "Authenticate exhibits and lay factual foundation.",
        "sections": ["caption", "declaration", "numbered paragraphs", "exhibit references", "signature", "proof of service"],
        "required_inputs": ["declarant identity", "facts to support", "exhibits"],
    },
    {
        "id": "proposed_order",
        "title": "Proposed Order",
        "category": "supporting",
        "purpose": "Provide judge-ready order language for motion disposition.",
        "sections": ["caption", "order title", "relief granted", "signature block"],
        "required_inputs": ["motion relief requested", "party names"],
    },
    {
        "id": "exhibit_index",
        "title": "Exhibit Index",
        "category": "supporting",
        "purpose": "Catalog exhibits, sources, and citations for filing or appendix use.",
        "sections": ["caption", "table of exhibits", "source citations", "authentication notes"],
        "required_inputs": ["exhibit list", "source citations", "page spans"],
    },
    {
        "id": "appendix",
        "title": "Appendix",
        "category": "appellate",
        "purpose": "Assemble record excerpts and documents for appellate filing.",
        "sections": ["caption", "appendix index", "record excerpts", "certificate of service"],
        "required_inputs": ["record items", "brief cites", "page limits"],
    },
    {
        "id": "timeline_summary",
        "title": "Timeline Summary",
        "category": "analysis",
        "purpose": "Compress record chronology into a filing-ready timeline.",
        "sections": ["caption", "timeline", "key events", "anomalies", "source anchors"],
        "required_inputs": ["event list", "dates", "source citations"],
    },
    {
        "id": "case_theory_memo",
        "title": "Case Theory Memo",
        "category": "analysis",
        "purpose": "State the strongest supported theory, alternatives, and pressure points.",
        "sections": ["caption", "question presented", "most supported theory", "alternate theories", "evidence supporting each", "risk notes"],
        "required_inputs": ["key facts", "top documents", "opposing facts"],
    },
    {
        "id": "filing_checklist",
        "title": "Filing Checklist",
        "category": "operations",
        "purpose": "Track rule, form, signature, service, and exhibit readiness.",
        "sections": ["caption", "checklist items", "court rule notes", "service", "final review"],
        "required_inputs": ["court profile", "filing type", "required attachments"],
    },
]


PII_PATTERNS: List[Dict[str, Any]] = [
    {"label": "email", "pattern": re.compile(r"(?i)\b[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}\b")},
    {"label": "ssn", "pattern": re.compile(r"\b\d{3}-\d{2}-\d{4}\b")},
    {"label": "phone", "pattern": re.compile(r"(?<!\d)(?:\+?1[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}(?!\d)")},
    {"label": "credit_card", "pattern": re.compile(r"\b(?:\d[ -]*?){13,16}\b")},
    {"label": "dob", "pattern": re.compile(r"(?i)\b(?:dob|date of birth)\s*[:\-]?\s*\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b")},
]


def _tokenize(text: str) -> List[str]:
    return re.findall(r"[a-z0-9]+", (text or "").lower())


def _top_terms(records: List[Dict[str, Any]], limit: int = 6) -> List[str]:
    counter: Counter[str] = Counter()
    for record in records:
        text = " ".join(
            [
                str(record.get("title", "")),
                str(record.get("text", "")),
                " ".join(record.get("tags", []) or []),
                json.dumps(record.get("metadata", {}), ensure_ascii=False),
            ]
        )
        counter.update(term for term in _tokenize(text) if len(term) > 3)
    return [term for term, _ in counter.most_common(limit)]


def list_court_profiles() -> List[Dict[str, Any]]:
    return COURT_PROFILES


def get_court_profile(profile_id: Optional[str]) -> Dict[str, Any]:
    profile_id = (profile_id or "federal_district_civil").strip()
    for profile in COURT_PROFILES:
        if profile["id"] == profile_id:
            return profile
    return COURT_PROFILES[0]


def list_filing_templates() -> List[Dict[str, Any]]:
    return FILING_TEMPLATES


def get_template(template_id: Optional[str]) -> Dict[str, Any]:
    template_id = (template_id or "case_theory_memo").strip()
    for template in FILING_TEMPLATES:
        if template["id"] == template_id:
            return template
    return FILING_TEMPLATES[0]


def court_profile_report(profile: Dict[str, Any]) -> Dict[str, Any]:
    required_keys = [
        "id",
        "name",
        "scope",
        "caption_notes",
        "motion_notes",
        "artifact_defaults",
        "local_rules_source",
        "template_priority",
        "page_limit_hint",
    ]
    missing = [key for key in required_keys if key not in profile or profile.get(key) in (None, "", [], {})]
    local_notes = profile.get("local_rules_notes") or []
    source_files = profile.get("source_files") or []
    if isinstance(local_notes, str):
        local_notes = [local_notes]
    if isinstance(source_files, str):
        source_files = [source_files]
    return {
        "profile": profile,
        "missing_fields": missing,
        "local_rules_notes": local_notes,
        "source_files": source_files,
        "notes": [
            "Treat this profile as a local drafting preset, not an authoritative substitute for current court rules.",
            "Verify page limits, judge-specific procedures, and filing format before export.",
            "Update this profile with court-specific observations as you validate matters.",
        ],
        "ready": not bool(missing),
    }


def scan_sensitive_text(text: str) -> List[Dict[str, Any]]:
    findings: List[Dict[str, Any]] = []
    sample = text or ""
    for item in PII_PATTERNS:
        matches = item["pattern"].finditer(sample)
        for match in matches:
            findings.append(
                {
                    "label": item["label"],
                    "value": match.group(0),
                    "start": match.start(),
                    "end": match.end(),
                }
            )
            if len(findings) >= 40:
                return findings
    return findings


def redact_sensitive_text(text: str) -> str:
    redacted = text or ""
    replacements = {
        "email": "[REDACTED EMAIL]",
        "ssn": "[REDACTED SSN]",
        "phone": "[REDACTED PHONE]",
        "credit_card": "[REDACTED CARD]",
        "dob": "[REDACTED DOB]",
    }
    for item in PII_PATTERNS:
        redacted = item["pattern"].sub(replacements[item["label"]], redacted)
    return redacted


def scan_packet(packet: Dict[str, Any]) -> Dict[str, Any]:
    sections = packet.get("sections", []) or []
    source_text = "\n\n".join(str(section) for section in sections)
    findings = scan_sensitive_text(source_text)
    return {
        "findings": findings,
        "count": len(findings),
        "labels": sorted({item["label"] for item in findings}),
        "has_sensitive_data": bool(findings),
    }


def default_matter(case_id: Optional[str], title: Optional[str] = None) -> Dict[str, Any]:
    matter_id = (case_id or "unassigned").strip() or "unassigned"
    matter_title = (title or matter_id.replace("-", " ").title()).strip()
    return {
        "case_id": matter_id,
        "title": matter_title,
        "court_profile_id": "federal_district_civil",
        "court_name": "Federal District Court",
        "district": "",
        "jurisdiction": "Federal",
        "matter_type": "civil",
        "practice_area": "Litigation",
        "plaintiff": "",
        "defendant": "",
        "counsel": "",
        "billing_increment_minutes": 15,
        "billing_rate": 0.0,
        "confidentiality_level": "Privileged",
        "notes": "",
        "updated_at": time.time(),
    }


def _dedupe_records(records: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    seen = set()
    result: List[Dict[str, Any]] = []
    for record in records:
        key = str(record.get("id") or record.get("citation") or record.get("text") or "")[:240]
        if key in seen:
            continue
        seen.add(key)
        result.append(record)
    return result


def rank_scenarios(records: List[Dict[str, Any]], query: str = "") -> List[Dict[str, Any]]:
    terms = _top_terms(records, limit=8)
    query_tokens = set(_tokenize(query))
    issue_map = {
        "discovery": "Discovery pressure points and missing responses are visible in the record.",
        "motion": "Motion posture looks likely based on the strongest evidence cluster.",
        "timeline": "The timeline is the main organizing axis and likely to drive the theory.",
        "contradiction": "The record contains conflicts that should be elevated as contradictions.",
        "billing": "Billing and staffing workflow appears to be relevant to operational follow-through.",
    }
    best_label = "record-backed theory"
    if {"compel", "discovery", "rfa", "rog", "rpd"} & query_tokens or "discovery" in terms:
        best_label = "discovery-pressure theory"
    elif {"dismiss", "jurisdiction", "standing"} & query_tokens:
        best_label = "threshold-defense theory"
    elif {"summary", "judgment", "undisputed"} & query_tokens:
        best_label = "record-cleanup / summary-judgment theory"
    elif {"appeal", "brief", "appendix"} & query_tokens:
        best_label = "appellate-record theory"
    theory = f"The strongest supported theory appears to be a {best_label} grounded in: {', '.join(terms) if terms else 'the current record'}."
    scenarios = [
        {
            "label": "Most supported scenario",
            "confidence": 0.82 if records else 0.35,
            "summary": theory,
        },
        {
            "label": "Alternative scenario",
            "confidence": 0.56,
            "summary": "A narrower reading of the record could support a different procedural or factual theory; verify the missing pieces before filing.",
        },
        {
            "label": "Adverse / cautionary scenario",
            "confidence": 0.44,
            "summary": "The record may also support the opposing side's framing, so the system should keep adverse facts visible and force human review.",
        },
    ]
    return scenarios


def detect_anomalies(records: List[Dict[str, Any]], timeline: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    anomalies: List[Dict[str, Any]] = []
    texts = " ".join(str(item.get("text", "")) for item in records).lower()
    timeline_text = " ".join(str(item.get("summary", "")) for item in timeline).lower()
    combined = texts + " " + timeline_text
    probes = [
        ("missing response", ["no response", "missing", "unanswered", "ignored"]),
        ("contradiction", ["contradiction", "inconsistent", "conflict", "dispute", "disagrees"]),
        ("deadline risk", ["deadline", "late", "untimely", "waive", "waiver"]),
        ("signature / authenticity", ["unsigned", "fake", "forged", "authenticity", "verify"]),
        ("metadata / chain-of-custody", ["metadata", "custody", "preserve", "upload path", "source"]),
    ]
    for label, words in probes:
        score = sum(1 for word in words if word in combined)
        if score:
            anomalies.append(
                {
                    "label": label,
                    "severity": min(1.0, 0.35 + score * 0.15),
                    "summary": f"Potential {label} issue found in the matter record; verify with source citations.",
                }
            )
    if not anomalies and records:
        anomalies.append(
            {
                "label": "record-density",
                "severity": 0.42,
                "summary": "The record is populated, but the system should still pressure-test for contradictory sources and missing exhibits.",
            }
        )
    return anomalies[:6]


def build_exhibit_index(records: List[Dict[str, Any]], limit: int = 12) -> List[Dict[str, Any]]:
    index: List[Dict[str, Any]] = []
    for i, record in enumerate(_dedupe_records(records)[:limit], start=1):
        index.append(
            {
                "label": f"Exhibit {chr(64 + ((i - 1) % 26) + 1)}",
                "title": record.get("title") or record.get("source_name") or record.get("id") or "Exhibit",
                "citation": record.get("citation") or record.get("source_name") or "",
                "page": record.get("metadata", {}).get("page") if isinstance(record.get("metadata", {}), dict) else None,
                "source_type": record.get("source_type") or "record",
                "summary": str(record.get("text", ""))[:180],
            }
        )
    return index


def estimate_billing(records: List[Dict[str, Any]], increment_minutes: int = 15, hourly_rate: float = 0.0) -> Dict[str, Any]:
    increment_minutes = max(1, int(increment_minutes or 15))
    base_minutes = 30 + (len(records) * 2)
    rounded_minutes = int(math.ceil(base_minutes / increment_minutes) * increment_minutes)
    hours = round(rounded_minutes / 60.0, 2)
    value = round(hours * float(hourly_rate or 0.0), 2)
    return {
        "increment_minutes": increment_minutes,
        "base_minutes": base_minutes,
        "rounded_minutes": rounded_minutes,
        "estimated_hours": hours,
        "estimated_value": value,
        "note": "Estimate is heuristic and should be edited by the lawyer or billing coordinator.",
    }


def build_filing_packet(
    *,
    template_id: str,
    matter: Dict[str, Any],
    records: List[Dict[str, Any]],
    timeline: List[Dict[str, Any]],
    query: str = "",
) -> Dict[str, Any]:
    template = get_template(template_id)
    profile = get_court_profile(matter.get("court_profile_id"))
    scenarios = rank_scenarios(records, query=query)
    anomalies = detect_anomalies(records, timeline)
    exhibit_index = build_exhibit_index(records, limit=12)
    top_records = _dedupe_records(records)[:6]
    source_notes = [
        f"[{item.get('citation') or item.get('source_name') or 'source'}] {str(item.get('text', ''))[:200]}"
        for item in top_records
    ]
    sections = []
    for section in template["sections"]:
        if section == "caption":
            sections.append(
                f"{matter.get('title', 'Matter')}\nCourt: {matter.get('court_name', 'Federal Court')} ({profile['name']})\nCase: {matter.get('case_id', 'unassigned')}"
            )
        elif section == "memorandum of points and authorities":
            sections.append(
                "Grounded memorandum placeholder:\n"
                + "\n".join(f"- {note}" for note in source_notes[:5])
                + ("\n- No grounded material loaded." if not source_notes else "")
            )
        elif section == "statement of facts":
            facts = "\n".join(f"- {note}" for note in source_notes[:6]) or "- No facts loaded."
            sections.append(facts)
        elif section == "argument":
            sections.append(
                "Argument scaffold:\n"
                "- State the requested relief.\n"
                "- Tie each element to a record citation.\n"
                "- Surface adverse facts and explain why they do not defeat the requested relief."
            )
        elif section in {"requested relief", "request for relief"}:
            sections.append("Requested relief: [insert relief tailored to the motion and court profile].")
        elif section == "declaration":
            sections.append("Declaration scaffold: authenticate the exhibits and establish foundation.")
        elif section == "exhibit index":
            sections.append(
                "Exhibit index:\n"
                + "\n".join(
                    f"- {item['label']}: {item['title']} ({item['citation'] or 'no citation'})"
                    for item in exhibit_index
                )
            )
        elif section == "proposed order":
            sections.append("Proposed order scaffold: state the order exactly as the court should enter it.")
        else:
            sections.append(f"{section.title()} scaffold: adapt to {profile['name']} and the selected matter.")
    return {
        "template": template,
        "court_profile": profile,
        "court_profile_report": court_profile_report(profile),
        "matter": matter,
        "records": records,
        "timeline": timeline,
        "scenarios": scenarios,
        "anomalies": anomalies,
        "exhibit_index": exhibit_index,
        "sections": sections,
        "sensitivity": scan_packet({"sections": sections}),
        "checklist": [
            "Confirm caption, docket, and parties",
            "Verify court and judge formatting rules",
            "Verify signatures and service requirements",
            "Cross-check citations against the loaded record",
            "Review for privilege, sealing, and redaction issues",
            "Have attorney review before filing",
        ],
    }


def packet_to_markdown(packet: Dict[str, Any], *, redact: bool = False) -> str:
    template = packet.get("template", {})
    matter = packet.get("matter", {})
    court_profile = packet.get("court_profile", {})
    court_profile_report_data = packet.get("court_profile_report", {})
    sensitivity = packet.get("sensitivity", {})
    scenarios = packet.get("scenarios", [])
    anomalies = packet.get("anomalies", [])
    exhibit_index = packet.get("exhibit_index", [])
    checklist = packet.get("checklist", [])
    sections = packet.get("sections", [])

    lines = [
        f"# {template.get('title', 'Filing Packet')}",
        "",
        "## Matter",
        f"- Case: {matter.get('title', 'Unassigned')}",
        f"- Case ID: {matter.get('case_id', 'unassigned')}",
        f"- Court: {matter.get('court_name', 'Federal Court')}",
        f"- Profile: {court_profile.get('name', 'Federal District Civil')}",
        f"- Practice Area: {matter.get('practice_area', 'Litigation')}",
        "",
        "## Court Profile Report",
        f"- Ready: {court_profile_report_data.get('ready', False)}",
        f"- Missing Fields: {', '.join(court_profile_report_data.get('missing_fields', [])) or 'none'}",
        f"- Local Rule Notes: {len(court_profile_report_data.get('local_rules_notes', []))}",
        f"- Source Files: {', '.join(court_profile_report_data.get('source_files', [])) or 'none'}",
        f"- Sensitive Findings: {sensitivity.get('count', 0)}",
        "",
        "## Scenario",
    ]
    if scenarios:
        for item in scenarios[:3]:
            lines.extend(
                [
                    f"- {item.get('label', 'Scenario')}: {item.get('summary', '')}",
                    f"  - Confidence: {item.get('confidence', 0)}",
                ]
            )
    else:
        lines.append("- No grounded scenario available.")

    lines.extend(["", "## Anomalies"])
    if anomalies:
        for item in anomalies:
            lines.append(f"- {item.get('label', 'anomaly')} ({item.get('severity', 0)}): {item.get('summary', '')}")
    else:
        lines.append("- No anomalies flagged.")

    lines.extend(["", "## Packet Sections"])
    for idx, section in enumerate(sections, start=1):
        lines.extend([f"### {idx}. {section.splitlines()[0][:120]}", section, ""])

    lines.extend(["## Exhibit Index"])
    if exhibit_index:
        for item in exhibit_index:
            lines.append(
                f"- {item.get('label', 'Exhibit')}: {item.get('title', 'Untitled')} | {item.get('citation', '')} | {item.get('summary', '')}"
            )
    else:
        lines.append("- No exhibits loaded.")

    lines.extend(["", "## Filing Checklist"])
    for item in checklist:
        lines.append(f"- {item}")

    markdown = "\n".join(lines).strip() + "\n"
    return redact_sensitive_text(markdown) if redact else markdown


def packet_to_docx_bytes(packet: Dict[str, Any], *, redact: bool = False) -> bytes:
    if docx is None:
        raise RuntimeError("python-docx is not available")

    template = packet.get("template", {})
    matter = packet.get("matter", {})
    court_profile = packet.get("court_profile", {})
    court_profile_report_data = packet.get("court_profile_report", {})
    sensitivity = packet.get("sensitivity", {})
    scenarios = packet.get("scenarios", [])
    anomalies = packet.get("anomalies", [])
    exhibit_index = packet.get("exhibit_index", [])
    checklist = packet.get("checklist", [])
    sections = packet.get("sections", [])

    document = docx.Document()
    try:
        document.styles["Normal"].font.name = "Aptos"
    except Exception:
        pass

    def add_bullets(lines: List[str]) -> None:
        for line in lines:
            document.add_paragraph(line, style="List Bullet")

    document.add_heading(template.get("title", "Filing Packet"), level=0)

    document.add_heading("Matter", level=1)
    add_bullets(
        [
            f"Case: {matter.get('title', 'Unassigned')}",
            f"Case ID: {matter.get('case_id', 'unassigned')}",
            f"Court: {matter.get('court_name', 'Federal Court')}",
            f"Profile: {court_profile.get('name', 'Federal District Civil')}",
            f"Practice Area: {matter.get('practice_area', 'Litigation')}",
        ]
    )

    document.add_heading("Court Profile Report", level=1)
    add_bullets(
        [
            f"Ready: {court_profile_report_data.get('ready', False)}",
            f"Missing Fields: {', '.join(court_profile_report_data.get('missing_fields', [])) or 'none'}",
            f"Local Rule Notes: {len(court_profile_report_data.get('local_rules_notes', []))}",
            f"Source Files: {', '.join(court_profile_report_data.get('source_files', [])) or 'none'}",
            f"Sensitive Findings: {sensitivity.get('count', 0)}",
        ]
    )

    document.add_heading("Scenario", level=1)
    if scenarios:
        for item in scenarios[:3]:
            document.add_paragraph(f"{item.get('label', 'Scenario')}: {item.get('summary', '')}", style="List Bullet")
            document.add_paragraph(f"Confidence: {item.get('confidence', 0)}")
    else:
        document.add_paragraph("No grounded scenario available.")

    document.add_heading("Anomalies", level=1)
    if anomalies:
        for item in anomalies:
            document.add_paragraph(
                f"{item.get('label', 'anomaly')} ({item.get('severity', 0)}): {item.get('summary', '')}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("No anomalies flagged.")

    document.add_heading("Packet Sections", level=1)
    for idx, section in enumerate(sections, start=1):
        heading = section.splitlines()[0][:120] if section else f"Section {idx}"
        document.add_heading(f"{idx}. {heading}", level=2)
        for paragraph in section.split("\n"):
            text = paragraph.strip()
            if text:
                document.add_paragraph(text)

    document.add_heading("Exhibit Index", level=1)
    if exhibit_index:
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        header = table.rows[0].cells
        header[0].text = "Label"
        header[1].text = "Title"
        header[2].text = "Citation"
        header[3].text = "Summary"
        for item in exhibit_index:
            row = table.add_row().cells
            row[0].text = str(item.get("label", "Exhibit"))
            row[1].text = str(item.get("title", "Untitled"))
            row[2].text = str(item.get("citation", ""))
            row[3].text = str(item.get("summary", ""))
    else:
        document.add_paragraph("No exhibits loaded.")

    document.add_heading("Filing Checklist", level=1)
    add_bullets([str(item) for item in checklist])

    if redact:
        for paragraph in document.paragraphs:
            paragraph.text = redact_sensitive_text(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = redact_sensitive_text(cell.text)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
